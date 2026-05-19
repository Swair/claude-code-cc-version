from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree
import re
import sys
import os

NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W = '{%s}' % NS_W
NS_XML = 'http://www.w3.org/XML/1998/namespace'
XML_SPACE = '{%s}space' % NS_XML

FONT_CN = 'SimSun'
FONT_EN = 'Times New Roman'
FONT_CODE = 'Consolas'

# Format profiles
FORMATS = {
    'cn': {
        'body_size': Pt(10.5),
        'heading1_size': Pt(14),
        'heading2_size': Pt(12),
        'heading3_size': Pt(10.5),
        'title_size': Pt(16),
        'author_size': Pt(12),
        'table_size': Pt(9),
        'caption_size': Pt(9),
        'line_spacing': Pt(16),
        'top_margin': Cm(2.54),
        'bottom_margin': Cm(2.54),
        'left_margin': Cm(3.17),
        'right_margin': Cm(3.17),
        'first_indent': Cm(0.74),
        'author_italic': True,
        'abs_indent': None,
        'table_caption': ('表', 'Tab.'),
        'figure_caption': ('图', 'Fig.'),
        'abs_label_cn': '摘要：',
        'abs_label_en': 'Abstract: ',
        'kw_label_cn': '关键词：',
        'kw_label_en': 'Keywords: ',
    },
    'en': {
        'body_size': Pt(11),
        'heading1_size': Pt(14),
        'heading2_size': Pt(12),
        'heading3_size': Pt(11),
        'title_size': Pt(18),
        'author_size': Pt(12),
        'table_size': Pt(9),
        'caption_size': Pt(10),
        'line_spacing': Pt(14.5),
        'top_margin': Cm(2.54),
        'bottom_margin': Cm(2.54),
        'left_margin': Cm(2.54),
        'right_margin': Cm(2.54),
        'first_indent': None,
        'author_italic': False,
        'abs_indent': None,
        'table_caption': ('Table', 'Table'),
        'figure_caption': ('Figure', 'Figure'),
        'abs_label_cn': 'Abstract: ',
        'abs_label_en': '',
        'abs_header_en': 'Abstract',
        'kw_label_cn': 'Keywords:',
        'kw_label_en': 'Keywords:',
    },
}


def _fmt(name, is_cn):
    p = 'cn' if is_cn else 'en'
    return FORMATS[p][name]


def _el(tag, **attrs):
    el = etree.Element(W + tag)
    for k, v in attrs.items():
        el.set(W + k, str(v))
    return el


def _sub(el, tag, **attrs):
    return etree.SubElement(el, W + tag, **{W + k: str(v) for k, v in attrs.items()})


def _rfonts(font_cn, font_en):
    return _el('rFonts', ascii=font_en, hAnsi=font_en, eastAsia=font_cn, cs=font_en)


def _set_run(run, font_cn, font_en, size, bold=False, italic=False, color=None, sup=False):
    run.font.size = size
    run.font.name = font_en
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rPr.append(_rfonts(font_cn, font_en))
    else:
        rf.set(qn('w:ascii'), font_en)
        rf.set(qn('w:hAnsi'), font_en)
        rf.set(qn('w:eastAsia'), font_cn)
        rf.set(qn('w:cs'), font_en)
    if sup:
        rPr.append(_el('vertAlign', val='superscript'))


def _pf(p, space_before=Pt(0), space_after=Pt(0), indent=None, align=None, is_cn=True):
    pf = p.paragraph_format
    pf.space_before = space_before
    pf.space_after = space_after
    if indent:
        pf.first_line_indent = indent
    if align:
        pf.alignment = align
    pf.line_spacing = _fmt('line_spacing', is_cn)


def _add_text(p, text, font_cn, font_en, size, bold=False, italic=False, color=None, sup=False):
    run = p.add_run(text)
    _set_run(run, font_cn, font_en, size, bold, italic, color, sup)
    return run


def _rich_text(p, text, font_cn, font_en, size, color=None, is_cn=True):
    """Parse **bold** and *italic* inline markdown."""
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if not part:
            continue
        bm = re.match(r'\*\*(.+?)\*\*', part)
        im = re.match(r'\*(.+?)\*', part)
        if bm:
            _add_text(p, bm.group(1), font_cn, font_en, size, bold=True, color=color)
        elif im:
            _add_text(p, im.group(1), font_cn, font_en, size, italic=True, color=color)
        else:
            refs = re.split(r'(\[\d+\])', part)
            for seg in refs:
                if re.match(r'^\[\d+\]$', seg):
                    _add_text(p, seg, FONT_EN, FONT_EN, size, color=RGBColor(0x00, 0x00, 0xFF), sup=True)
                elif seg:
                    _add_text(p, seg, font_cn, font_en, size, color=color)


def md_to_paper(md_file, docx_file, generate_pdf=False):
    with open(md_file, 'r', encoding='utf-8-sig') as f:
        content = f.read()

    is_cn = bool(re.search(r'[一-鿿]', md_file))
    fmt = FORMATS['cn' if is_cn else 'en']

    # --- Strip Software Citation / 软件引用 section and footer ---
    content = re.split(r'##\s+(?:Software Citation|软件引用)\b', content, flags=re.IGNORECASE)[0]
    content = re.sub(r'\*Version:.*\*\*\s*\n.*\*Date:.*\n.*\*Author:.*\n?', '', content)
    content = content.rstrip('\n') + '\n'

    # --- Extract YAML frontmatter metadata ---
    # Only parse YAML if content starts with --- (actual YAML frontmatter)
    meta = {}
    if content.startswith('---'):
        parts = content.split('---', 2)
        yaml_block = parts[1] if len(parts) > 1 else ''
        # Verify it's actual YAML (has key: value pairs) before treating as metadata
        has_yaml = any(': ' in ml.strip() or ml.strip().endswith(':') for ml in yaml_block.split('\n') if ml.strip())
        if has_yaml:
            content = parts[2].lstrip('\n') if len(parts) > 2 else content
            for ml in yaml_block.split('\n'):
                ml = ml.strip()
                if ':' in ml:
                    k, v = ml.split(':', 1)
                    meta[k.strip()] = v.strip().strip('"').strip("'")

    title_line = ''
    title_match = re.match(r'^#\s+(.+)', content)
    if title_match:
        title_line = title_match.group(1).strip()
        # Remove document-type prefix (e.g. "Whitepaper: ", "论文: ")
        title_line = re.sub(r'^(?:Whitepaper|论文|白皮书)\s*[：:]\s*', '', title_line)
        rest = content[title_match.end():].lstrip('\n')
        # Strip ALL blockquote lines right after title (metadata, notes, etc.)
        while rest.startswith('>'):
            rest = re.sub(r'^>[^\n]*\n?', '', rest, flags=re.MULTILINE)
            rest = rest.lstrip('\n')
        # Skip author plain-text line (from YAML or standalone)
        author = meta.get('author', '')
        if author:
            if rest.startswith(author):
                rest = rest[len(author):].lstrip('\n')
        else:
            # Fallback: detect author name from plain-text line
            first_line_m = re.match(r'^([^\n#|>]+)\n', rest)
            if first_line_m:
                candidate = first_line_m.group(1).strip()
                if candidate and len(candidate) < 50 and not candidate.startswith('#'):
                    author = candidate
                    rest = rest[first_line_m.end():].lstrip('\n')
        # Store extracted author so rendering code can use it
        meta['author'] = author
        content = rest

    doc = Document()

    for sec in doc.sections:
        sec.top_margin = fmt['top_margin']
        sec.bottom_margin = fmt['bottom_margin']
        sec.left_margin = fmt['left_margin']
        sec.right_margin = fmt['right_margin']

    # --- Page number in footer ---
    footer = sec.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = fp._element.get_or_add_pPr()
    tabs = _sub(pPr, 'tabs')
    _sub(tabs, 'tab', val='center', pos='4680')

    run = fp.add_run()
    _sub(run._element, 'fldChar', fldCharType='begin')
    _sub(run._element, 'instrText', xml__space='preserve').text = ' PAGE '
    run2 = fp.add_run()
    _sub(run2._element, 'fldChar', fldCharType='end')
    _set_run(run, FONT_EN, FONT_EN, Pt(9))
    _set_run(run2, FONT_EN, FONT_EN, Pt(9))

    # --- Title ---
    if title_line:
        p = doc.add_paragraph()
        if is_cn:
            _add_text(p, title_line, FONT_CN, FONT_EN, fmt['title_size'], bold=True)
        else:
            _add_text(p, title_line, FONT_EN, FONT_EN, fmt['title_size'], bold=True)
        _pf(p, Pt(12), Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER, is_cn=is_cn)

    # --- Class number / document code (Chinese format only) ---
    if is_cn:
        cn = meta.get('zhongtu_fenhao', '') or meta.get('classification', '')
        if cn:
            p = doc.add_paragraph()
            _add_text(p, '中图分类号: ' + cn, FONT_CN, FONT_EN, Pt(9))
            _pf(p, Pt(2), Pt(2), is_cn=True)
        dm = meta.get('wenxian_mark', '') or meta.get('document_code', '') or meta.get('标志码', '')
        if dm:
            p = doc.add_paragraph()
            _add_text(p, '文献标志码: ' + dm, FONT_CN, FONT_EN, Pt(9))
            _pf(p, Pt(2), Pt(2), is_cn=True)

    # --- Author ---
    author = meta.get('author', '')
    if author:
        p = doc.add_paragraph()
        if is_cn:
            _add_text(p, author, FONT_CN, FONT_EN, fmt['author_size'])
        else:
            _add_text(p, author, FONT_EN, FONT_EN, fmt['author_size'])
        _pf(p, Pt(6), Pt(10), align=WD_ALIGN_PARAGRAPH.CENTER, is_cn=is_cn)

    # --- Normal style ---
    st = doc.styles['Normal']
    st.font.size = fmt['body_size']
    st.font.name = FONT_EN
    st.paragraph_format.space_before = Pt(0)
    st.paragraph_format.space_after = Pt(0)
    st.paragraph_format.line_spacing = fmt['line_spacing']
    st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)

    for pn in ['Heading1', 'Heading2', 'Heading3']:
        if pn in [s.name for s in doc.styles]:
            s = doc.styles[pn]
            s.font.name = FONT_EN
            s.font.bold = True
            s.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)

    lines = content.split('\n')
    in_code = False
    code_buf = []
    h1 = h2 = h3 = 0
    table_count = 0
    figure_count = 0
    i = 0

    # Font selection based on language
    body_fc_cn = FONT_CN
    body_fc_en = FONT_EN

    while i < len(lines):
        line = lines[i]

        # --- Fenced code blocks ---
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                p = doc.add_paragraph()
                _add_text(p, '\n'.join(code_buf), FONT_CODE, FONT_CODE, Pt(9))
                _pf(p, Pt(6), Pt(6), is_cn=is_cn)
                pPr = p._element.get_or_add_pPr()
                pPr.append(_el('shd', val='clear', color='none', fill='F5F5F5'))
                pBdr = etree.Element(W + 'pBdr')
                for side in ['top', 'bottom', 'left', 'right']:
                    c = _sub(pBdr, side, val='single', sz='4', space='4', color='CCCCCC')
                pPr.append(pBdr)
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # --- Abstract / Keywords special handling ---
        hm_abstract = re.match(r'^(#{1,6})\s+(.*)', line)
        if hm_abstract:
            heading_text = hm_abstract.group(2).strip()
            is_abs_heading = re.search(r'摘要|\bAbstract\b', heading_text) and not re.search(r'关键词|Keywords?', heading_text)
            if is_abs_heading:
                abs_lines = []
                i += 1
                while i < len(lines):
                    nl = lines[i]
                    nhm = re.match(r'^(#{1,6})\s+(.*)', nl)
                    if nhm and not re.search(r'摘要|\bAbstract\b', nhm.group(2)):
                        break
                    if re.match(r'^-{3,}$', nl.strip()) or re.match(r'^\*{3,}$', nl.strip()):
                        i += 1
                        continue
                    if nl.strip() == '' and abs_lines and abs_lines[-1].strip() == '':
                        i += 1
                        continue
                    abs_lines.append(nl)
                    i += 1

                kw_idx = None
                for ki in range(len(abs_lines) - 1, -1, -1):
                    if re.search(r'关键词|Keywords?', abs_lines[ki]):
                        kw_idx = ki
                        break

                abs_body_lines = (abs_lines[:kw_idx] if kw_idx is not None else abs_lines)
                kw_text = abs_lines[kw_idx].strip() if kw_idx is not None else ''
                fc = FONT_CN if is_cn else FONT_EN
                pf_indent = fmt['abs_indent']

                # Split abs_body_lines into blocks: tables and text paragraphs
                def _render_abs_table(rows):
                    nc = max(len(r) for r in rows)
                    t = doc.add_table(rows=len(rows), cols=nc)
                    t.style = 'Table Grid'
                    tbl = t._element
                    tblPr = tbl.find(qn('w:tblPr'))
                    if tblPr is None:
                        tblPr = _el('tblPr')
                        _sub(tblPr, 'tblW', w='5000', type='dxa')
                        tbl.insert(0, tblPr)
                    for ri, row in enumerate(rows):
                        for ci in range(nc):
                            cell = t.cell(ri, ci)
                            cell.text = row[ci] if ci < len(row) else ''
                            for par in cell.paragraphs:
                                _pf(par, Pt(2), Pt(2), is_cn=is_cn)
                                for run in par.runs:
                                    _set_run(run, FONT_CN, FONT_EN, fmt['table_size'])
                                if ri == 0:
                                    for run in par.runs:
                                        run.bold = True
                                        _set_run(run, FONT_CN, FONT_EN, fmt['table_size'], bold=True)
                                    tc = cell._element
                                    tcPr = tc.get_or_add_tcPr()
                                    tcPr.append(_el('shd', val='clear', color='none', fill='D9E2F3'))
                                    par.alignment = WD_ALIGN_PARAGRAPH.CENTER

                def _render_abs_text(text, is_first, label_text):
                    p = doc.add_paragraph()
                    if is_first and label_text:
                        _add_text(p, label_text, fc, FONT_EN, fmt['body_size'], bold=True)
                    _rich_text(p, text, fc, FONT_EN, fmt['body_size'], is_cn=is_cn)
                    _pf(p, Pt(6), Pt(4), indent=pf_indent, is_cn=is_cn)

                # Parse blocks: group text into paragraphs, render tables separately
                bi = 0
                first_para = True
                label_text = fmt['abs_label_cn'] if is_cn else fmt['abs_label_en']
                # Standalone ABSTRACT header for arXiv style
                abs_header = fmt.get('abs_header_en', '')
                if abs_header and not is_cn:
                    ph = doc.add_paragraph()
                    _add_text(ph, abs_header, FONT_EN, FONT_EN, fmt['heading1_size'], bold=True)
                    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _pf(ph, Pt(8), Pt(8), is_cn=is_cn)
                text_buf = []
                def _flush_text():
                    nonlocal first_para, label_text
                    if not text_buf:
                        return
                    joined = ' '.join(text_buf)
                    p = doc.add_paragraph()
                    if first_para and label_text:
                        _add_text(p, label_text, fc, FONT_EN, fmt['body_size'], bold=True)
                        first_para = False
                        label_text = ''
                    _rich_text(p, joined, fc, FONT_EN, fmt['body_size'], is_cn=is_cn)
                    _pf(p, Pt(6), Pt(4), indent=pf_indent, is_cn=is_cn)
                    text_buf[:] = []
                while bi < len(abs_body_lines):
                    bl = abs_body_lines[bi]
                    # Check for table
                    if '|' in bl and bi + 1 < len(abs_body_lines) and re.match(r'^[\s|:-]+$', abs_body_lines[bi + 1]):
                        _flush_text()
                        tbl_rows = []
                        sep_seen = False
                        while bi < len(abs_body_lines):
                            if '|' in abs_body_lines[bi]:
                                if not sep_seen and re.match(r'^[\s|:-]+$', abs_body_lines[bi].strip()):
                                    sep_seen = True
                                    bi += 1
                                    continue
                                cells = [c.strip() for c in abs_body_lines[bi].split('|')]
                                if cells and cells[0] == '': cells = cells[1:]
                                if cells and cells[-1] == '': cells = cells[:-1]
                                if cells: tbl_rows.append(cells)
                                bi += 1
                            else:
                                break
                        if tbl_rows:
                            _render_abs_table(tbl_rows)
                        continue
                    # Text line (non-empty)
                    stripped = bl.strip()
                    if stripped and not re.match(r'^-{3,}$', stripped):
                        text_buf.append(stripped)
                    elif not stripped:
                        _flush_text()
                    bi += 1
                _flush_text()

                # Render keywords
                if kw_text:
                    p = doc.add_paragraph()
                    kw_clean = kw_text
                    kw_clean = re.sub(r'^\*\*(?:关键词|Keywords?)\*\*\s*[：:]\s*', '', kw_clean)
                    kw_clean = re.sub(r'^(?:关键词|Keywords?)[：:]\s*', '', kw_clean)
                    kw_clean = kw_clean.lstrip(':')
                    _add_text(p, fmt['kw_label_cn'] if is_cn else fmt['kw_label_en'], fc, FONT_EN, fmt['body_size'], bold=True)
                    if kw_clean:
                        _add_text(p, kw_clean, FONT_CN, FONT_EN, fmt['body_size'])
                    _pf(p, Pt(2), Pt(8), indent=pf_indent, is_cn=is_cn)
                continue

        # --- Headings ---
        hm = re.match(r'^(#{1,6})\s+(.*)', line)
        if hm:
            level = len(hm.group(1))
            text = hm.group(2).strip().rstrip('#').strip()
            fc = FONT_CN if is_cn else FONT_EN

            if level == 1:
                p = doc.add_paragraph()
                _add_text(p, text, fc, FONT_EN, fmt['heading1_size'], bold=True)
                _pf(p, Pt(12), Pt(6), is_cn=is_cn)
            elif level == 2:
                p = doc.add_paragraph()
                _add_text(p, text, fc, FONT_EN, fmt['heading2_size'], bold=True)
                _pf(p, Pt(8), Pt(4), is_cn=is_cn)
            elif level == 3:
                p = doc.add_paragraph()
                _add_text(p, text, fc, FONT_EN, fmt['heading3_size'], bold=True)
                _pf(p, Pt(6), Pt(3), is_cn=is_cn)
            elif level == 4:
                p = doc.add_paragraph()
                _add_text(p, text, fc, FONT_EN, fmt['body_size'], bold=True)
                _pf(p, Pt(4), Pt(2), is_cn=is_cn)
            else:
                p = doc.add_paragraph()
                _add_text(p, text, fc, FONT_EN, fmt['body_size'], bold=True)
                _pf(p, Pt(3), Pt(2), is_cn=is_cn)
            i += 1
            continue

        # --- Horizontal rule ---
        if re.match(r'^-{3,}$', line.strip()) or re.match(r'^\*{3,}$', line.strip()):
            p = doc.add_paragraph()
            _pf(p, Pt(8), Pt(8), is_cn=is_cn)
            pPr = p._element.get_or_add_pPr()
            pPr.append(_el('pBdr', bottom_val='single', bottom_sz='6', bottom_space='1', bottom_color='AAAAAA'))
            i += 1
            continue

        # --- Table ---
        if '|' in line and i + 1 < len(lines) and re.match(r'^[\s|:-]+$', lines[i + 1]):
            rows = []
            sep_seen = False
            while i < len(lines):
                if '|' in lines[i]:
                    if not sep_seen and re.match(r'^[\s|:-]+$', lines[i].strip()):
                        sep_seen = True
                        i += 1
                        continue
                    cells = [c.strip() for c in lines[i].split('|')]
                    if cells and cells[0] == '': cells = cells[1:]
                    if cells and cells[-1] == '': cells = cells[:-1]
                    if cells: rows.append(cells)
                    i += 1
                else:
                    break

            if rows:
                table_count += 1
                nc = max(len(r) for r in rows)

                # Caption ABOVE table
                pc = doc.add_paragraph()
                cap_label = fmt['table_caption'][0] if is_cn else fmt['table_caption'][1]
                fc = FONT_CN if is_cn else FONT_EN
                _add_text(pc, cap_label + ' ' + str(table_count) + '  ', fc, FONT_EN, fmt['caption_size'])
                pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _pf(pc, Pt(2), Pt(4), is_cn=is_cn)

                # Table
                t = doc.add_table(rows=len(rows), cols=nc)
                t.style = 'Table Grid'
                tbl = t._element
                tblPr = tbl.find(qn('w:tblPr'))
                if tblPr is None:
                    tblPr = _el('tblPr')
                    _sub(tblPr, 'tblW', w='5000', type='dxa')
                    tbl.insert(0, tblPr)

                for ri, row in enumerate(rows):
                    for ci in range(nc):
                        cell = t.cell(ri, ci)
                        cell.text = row[ci] if ci < len(row) else ''
                        for par in cell.paragraphs:
                            _pf(par, Pt(2), Pt(2), is_cn=is_cn)
                            for run in par.runs:
                                _set_run(run, FONT_CN, FONT_EN, fmt['table_size'])
                            if ri == 0:
                                for run in par.runs:
                                    run.bold = True
                                    _set_run(run, FONT_CN, FONT_EN, fmt['table_size'], bold=True)
                                tc = cell._element
                                tcPr = tc.get_or_add_tcPr()
                                tcPr.append(_el('shd', val='clear', color='none', fill='D9E2F3'))
                                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # --- Figure (markdown image) ---
        im = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)', line.strip())
        if im:
            alt = im.group(1).strip()
            img_path = im.group(2).strip()
            figure_count += 1
            full = img_path
            if not os.path.isabs(img_path):
                full = os.path.join(os.path.dirname(os.path.abspath(md_file)), img_path)
            if os.path.exists(full):
                p = doc.add_paragraph()
                p.add_run().add_picture(full)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _pf(p, Pt(6), Pt(2), is_cn=is_cn)
            # Caption below figure
            pc = doc.add_paragraph()
            cap_label = fmt['figure_caption'][0] if is_cn else fmt['figure_caption'][1]
            fc = FONT_CN if is_cn else FONT_EN
            _add_text(pc, cap_label + ' ' + str(figure_count), fc, FONT_EN, fmt['caption_size'])
            if alt:
                _add_text(pc, '  ' + alt, FONT_CN, FONT_EN, fmt['caption_size'])
            pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _pf(pc, Pt(2), Pt(4), is_cn=is_cn)
            i += 1
            continue

        # --- Unordered list ---
        um = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if um:
            lvl = len(um.group(1)) // 2
            text = um.group(2).strip()
            p = doc.add_paragraph()
            _add_text(p, '  ' * lvl + '• ' + text, FONT_CN, FONT_EN, fmt['body_size'])
            _pf(p, Pt(3), Pt(3), is_cn=is_cn)
            i += 1
            continue

        # --- Ordered list ---
        om = re.match(r'^(\s*)\d+[.)]\s+(.*)', line)
        if om:
            text = om.group(2).strip()
            p = doc.add_paragraph()
            _add_text(p, text, FONT_CN, FONT_EN, fmt['body_size'])
            _pf(p, Pt(3), Pt(3), indent=Cm(0.74) if is_cn else None, is_cn=is_cn)
            i += 1
            continue

        # --- Empty line ---
        if line.strip() == '':
            i += 1
            continue

        # --- Blockquote ---
        bm = re.match(r'^>\s?(.*)', line)
        if bm:
            text = bm.group(1).strip()
            # Standalone '>' with no content - skip like empty line
            if not text:
                i += 1
                continue
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            pBdr = _el('pBdr')
            _sub(pBdr, 'left', val='single', sz='4', space='4', color='6699CC')
            pPr.append(pBdr)
            fc = FONT_CN if is_cn else FONT_EN
            _rich_text(p, text, fc, FONT_EN, fmt['body_size'], color=RGBColor(0x66, 0x66, 0x66), is_cn=is_cn)
            _pf(p, Pt(6), Pt(6), is_cn=is_cn)
            for run in p.runs:
                run.font.italic = True
            i += 1
            continue

        # --- Regular paragraph ---
        p = doc.add_paragraph()
        fc = FONT_CN if is_cn else FONT_EN
        _rich_text(p, line.strip(), fc, FONT_EN, fmt['body_size'], is_cn=is_cn)
        _pf(p, indent=fmt['first_indent'], is_cn=is_cn)
        i += 1

    doc.save(docx_file)
    print(f'OK: {docx_file}')

    if generate_pdf:
        _generate_pdf(docx_file, docx_file.replace('.docx', '.pdf'))

def _generate_pdf(docx_path, pdf_path):
    """Convert docx to PDF using Word COM automation (Windows)."""
    try:
        from docx2pdf import convert
        import os
        convert(docx_path, pdf_path)
        base = os.path.basename(pdf_path)
        print(f'OK: {pdf_path}')
    except Exception as e:
        print(f'PDF WARN: {e}')

if __name__ == '__main__':
    md_file = sys.argv[1]
    docx_file = sys.argv[2] if len(sys.argv) > 2 else ''
    generate_pdf = '--pdf' in sys.argv
    if not docx_file:
        base = os.path.splitext(os.path.basename(md_file))[0]
        docx_file = f'skills_output/{base}.docx'
    md_to_paper(md_file, docx_file, generate_pdf)
